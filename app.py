import streamlit as st
import pandas as pd
import altair as alt
import re
import openpyxl
from io import BytesIO
from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

# --- FONT CONSTANTS ---
BODY_FONT = "Verdana"
BODY_SIZE = 9
HEADER_FONT = "Verdana"
HEADER_SIZE = 10

# Matches a line that starts with manual numbering, e.g. "1. ", "2) ", "(3) ",
# nested numbering like "1.2. " / "2.3) ", and also "1.Text" with no space
# after the marker (a common typo in pasted Excel content). The (?!\d) guard
# stops a plain decimal like "9.8 is the CVSS score" from being misread as
# marker "9." followed by "8 is...".
NUMBERED_LINE_RE = re.compile(r'^\s*\(?\d+(?:\.\d+)*[\.\)](?!\d)\s*')

# --- RISK BADGE COLOR SCHEME (hex, no '#') ---
RISK_COLORS = {
    "CRITICAL": {"bg": "C00000", "text": "FFFFFF"},
    "HIGH":     {"bg": "EE0000", "text": "FFFFFF"},
    "MEDIUM":   {"bg": "FFC000", "text": "000000"},
    "LOW":      {"bg": "92D050", "text": "000000"},
    "INFO":     {"bg": "00B0F0", "text": "000000"},
}
FALLBACK_RISK_COLOR = {"bg": "D9D9D9", "text": "000000"}

SEVERITY_ORDER = {"CRITICAL": 0, "HIGH": 1, "MEDIUM": 2, "LOW": 3, "INFO": 4}


def normalize_risk_key(risk_text):
    key = str(risk_text).strip().upper()
    if key.startswith("INFORM"):
        key = "INFO"
    return key


# Column widths for 5-column summary tables (total = 6.5" for 1" margins)
SUMMARY_TABLE_COL_WIDTHS = [Inches(0.9), Inches(0.9), Inches(3.0), Inches(1.0), Inches(0.7)]

# Column widths for 3-column Host Config summary table (total = 6.5")
HOST_CONFIG_TABLE_COL_WIDTHS = [Inches(1.2), Inches(2.9), Inches(2.4)]


# --- REPORT TYPE CONFIGURATIONS ---

# Summary table: (header_label, excel_col) for columns 2–5 (col 1 is always Risk Rating)
REPORT_SUMMARY_FIELDS = {
    "Web Application VAPT": [
        ("Overall Score", "CVSS Score"),
        ("CVSS Vector",   "CVSS Vector"),
        ("OWASP Top 10",  "OWASP Top 10"),
        ("CWE ID",        "CWE ID"),
    ],
    "Network VAPT": [
        ("Overall Score", "CVSS Score"),
        ("CVSS Vector",   "CVSS Vector"),
        ("CVE ID",        "CVE ID"),
        ("Port/Protocol", "Port/Protocol"),
    ],
    "Cloud VAPT": [
        ("Overall Score",   "CVSS Score"),
        ("CVSS Vector",     "CVSS Vector"),
        ("Cloud Service",   "Cloud Service"),
        ("Resource ID/ARN", "Resource ID/ARN"),
    ],
}

# (display label, excel column name) for the affected asset field per report type
AFFECTED_ASSET_CONFIG = {
    "Web Application VAPT":    ("Affected Module(s)",     "Affected Module / URL"),
    "Network VAPT":             ("Affected Host/IP",       "Affected Host/IP"),
    "Host Configuration Review":("Affected Host/System",   "Affected Host/System"),
    "Cloud VAPT":               ("Affected Cloud Resource", "Affected Cloud Resource"),
}

# Per-type editable metadata fields in the UI: (label, session_key_prefix, excel_col)
# Split into two groups of 2 for the c2/c3 columns layout; Host Config only has 2 total.
REPORT_EDIT_FIELDS = {
    "Web Application VAPT": [
        ("CVSS Score",   "edit_f1", "CVSS Score"),
        ("CVSS Vector",  "edit_f2", "CVSS Vector"),
        ("OWASP Top 10", "edit_f3", "OWASP Top 10"),
        ("CWE ID",       "edit_f4", "CWE ID"),
    ],
    "Network VAPT": [
        ("CVSS Score",    "edit_f1", "CVSS Score"),
        ("CVSS Vector",   "edit_f2", "CVSS Vector"),
        ("CVE ID",        "edit_f3", "CVE ID"),
        ("Port/Protocol", "edit_f4", "Port/Protocol"),
    ],
    "Host Configuration Review": [
        ("Benchmark/Standard", "edit_f1", "Benchmark/Standard"),
        ("Section/Control",    "edit_f2", "Section/Control"),
    ],
    "Cloud VAPT": [
        ("CVSS Score",      "edit_f1", "CVSS Score"),
        ("CVSS Vector",     "edit_f2", "CVSS Vector"),
        ("Cloud Service",   "edit_f3", "Cloud Service"),
        ("Resource ID/ARN", "edit_f4", "Resource ID/ARN"),
    ],
}

# Finding body sections — same structure for all report types
FINDING_SECTIONS = [
    ("Observation",         "Observations",             True),
    ("Implication",         "Implications",             True),
    ("Recommendations",     "Recommendations",          True),
    ("Management Comments", "Management Comments",      False),
    ("Follow-up Comments",  "Post Review Observations", True),
    ("Status",              "Status",                   False),
]


# --- HELPER FUNCTIONS FOR FONT / TABLE STYLING ---

def apply_font(font, name=BODY_FONT, size=BODY_SIZE, bold=None, color=None):
    """
    Force a font onto a python-docx Font object.
    Strips theme font attributes so the override always takes effect in Word.
    `color`, if given, is a hex string without '#' (e.g. "000000").
    """
    font.name = name
    font.size = Pt(size)
    if bold is not None:
        font.bold = bold
    if color is not None:
        font.color.rgb = RGBColor.from_string(color)

    rPr = font._element.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rPr.append(rFonts)

    for attr in ('w:ascii', 'w:hAnsi', 'w:eastAsia', 'w:cs'):
        rFonts.set(qn(attr), name)
    for theme_attr in ('w:asciiTheme', 'w:hAnsiTheme', 'w:eastAsiaTheme', 'w:cstheme'):
        if rFonts.get(qn(theme_attr)) is not None:
            del rFonts.attrib[qn(theme_attr)]


def set_cell_background(cell, fill_color):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), fill_color)
    tcPr.append(shd)


def set_cell_border(cell, **kwargs):
    """
    Set borders on a table cell.
    Usage: set_cell_border(cell, top={"sz": 24, "val": "single", "color": "000000"})
    `sz` is in eighths of a point (e.g. 24 = 3pt, 6 = 0.75pt).
    """
    tcPr = cell._tc.get_or_add_tcPr()
    tcBorders = tcPr.find(qn('w:tcBorders'))
    if tcBorders is None:
        tcBorders = OxmlElement('w:tcBorders')
        tcPr.append(tcBorders)
    for edge in ('top', 'left', 'bottom', 'right'):
        edge_data = kwargs.get(edge)
        if not edge_data:
            continue
        tag = qn(f'w:{edge}')
        element = tcBorders.find(tag)
        if element is None:
            element = OxmlElement(f'w:{edge}')
            tcBorders.append(element)
        for key in ('sz', 'val', 'color', 'space'):
            if key in edge_data:
                element.set(qn(f'w:{key}'), str(edge_data[key]))


def get_risk_style(risk_text):
    return RISK_COLORS.get(normalize_risk_key(risk_text), FALLBACK_RISK_COLOR)


def sort_by_severity(df, risk_col):
    if risk_col not in df.columns:
        return df
    rank = df[risk_col].apply(lambda r: SEVERITY_ORDER.get(normalize_risk_key(r), len(SEVERITY_ORDER)))
    return (
        df.assign(_severity_rank=rank)
          .sort_values('_severity_rank', kind='stable')
          .drop(columns='_severity_rank')
          .reset_index(drop=True)
    )


def set_table_column_widths(table, widths):
    """python-docx needs width set on the table, AND on every cell in each
    column, AND autofit disabled — otherwise Word ignores it."""
    table.autofit = False
    table.allow_autofit = False
    for row in table.rows:
        for idx, width in enumerate(widths):
            row.cells[idx].width = width
    for idx, width in enumerate(widths):
        table.columns[idx].width = width


def _build_summary_table(doc, headers, values, col_widths):
    """
    Generic bordered summary table. values[0] is the risk rating and drives
    the badge color on the first data cell.
    """
    risk = values[0]
    style = get_risk_style(risk)
    n_cols = len(headers)

    table = doc.add_table(rows=2, cols=n_cols)
    set_table_column_widths(table, col_widths)

    # Header row
    for col, text in enumerate(headers):
        cell = table.cell(0, col)
        cell.text = text
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        para = cell.paragraphs[0]
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        para.paragraph_format.space_before = Pt(6)
        para.paragraph_format.space_after = Pt(6)
        for run in para.runs:
            apply_font(run.font, BODY_FONT, BODY_SIZE, bold=True)
        set_cell_border(
            cell,
            top={"sz": 24, "val": "single", "color": "000000"},
            bottom={"sz": 6, "val": "single", "color": "000000"},
        )

    # Data row
    for col, text in enumerate(values):
        cell = table.cell(1, col)
        cell.text = text
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        para = cell.paragraphs[0]
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        para.paragraph_format.space_before = Pt(6)
        para.paragraph_format.space_after = Pt(6)
        is_risk_cell = (col == 0)
        for run in para.runs:
            apply_font(run.font, BODY_FONT, BODY_SIZE, bold=is_risk_cell)
            if is_risk_cell:
                run.font.color.rgb = RGBColor.from_string(style["text"])
        set_cell_background(cell, style["bg"] if is_risk_cell else "F2F2F2")
        set_cell_border(cell, bottom={"sz": 6, "val": "single", "color": "000000"})

    return table


def add_five_col_summary_table(doc, risk, score, vector, field4, field5,
                                field4_label, field5_label):
    """5-column summary table used by Web VAPT, Network VAPT, and Cloud VAPT."""
    headers = ["Risk Rating", "Overall Score", "CVSS Vector", field4_label, field5_label]
    values  = [str(risk).upper(), str(score), str(vector), str(field4), str(field5)]
    return _build_summary_table(doc, headers, values, SUMMARY_TABLE_COL_WIDTHS)


def add_host_config_summary_table(doc, risk, benchmark, section_control):
    """3-column summary table for Host Configuration Review (no CVSS columns)."""
    headers = ["Risk Rating", "Benchmark/Standard", "Section/Control"]
    values  = [str(risk).upper(), str(benchmark), str(section_control)]
    return _build_summary_table(doc, headers, values, HOST_CONFIG_TABLE_COL_WIDTHS)


def add_styled_heading(doc, text, level):
    heading = doc.add_heading(text, level)
    for run in heading.runs:
        apply_font(run.font, HEADER_FONT, HEADER_SIZE, color="000000")
    if level == 3:
        heading.paragraph_format.space_after = Pt(6)
    return heading


def add_text_block(doc, value, justify=False):
    """
    Writes cell content as one or more paragraphs. Lines that look like
    manually-numbered items get hanging-indent list formatting.
    """
    if value is None or (isinstance(value, float) and pd.isna(value)):
        text = "N/A"
    else:
        text = str(value)

    lines = [ln.strip() for ln in text.splitlines() if ln.strip()]
    if not lines:
        lines = [text]

    for line in lines:
        match = NUMBERED_LINE_RE.match(line)
        if match:
            marker = match.group(0).strip()
            rest = line[match.end():].strip()
            line = f"{marker} {rest}" if rest else marker

        p = doc.add_paragraph(line)
        if justify:
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        if match:
            pf = p.paragraph_format
            pf.left_indent = Cm(0.63)
            pf.first_line_indent = -Cm(0.63)
            pf.space_before = Pt(0)
            pf.space_after = Pt(10)
            pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
            pf.line_spacing = 1.15
        for run in p.runs:
            apply_font(run.font, BODY_FONT, BODY_SIZE)


def add_styled_paragraph(doc, text="", style=None, alignment=None):
    p = doc.add_paragraph(text, style=style) if style else doc.add_paragraph(text)
    if alignment is not None:
        p.alignment = alignment
    for run in p.runs:
        apply_font(run.font, BODY_FONT, BODY_SIZE)
    return p


# --- SHARED DOCUMENT HELPERS ---

def _init_doc_styles(doc):
    apply_font(doc.styles['Normal'].font, BODY_FONT, BODY_SIZE)
    apply_font(doc.styles['Title'].font, HEADER_FONT, HEADER_SIZE, color="000000")
    for i in range(1, 4):
        apply_font(doc.styles[f'Heading {i}'].font, HEADER_FONT, HEADER_SIZE, color="000000")
    doc.styles['Heading 3'].paragraph_format.space_after = Pt(6)
    apply_font(doc.styles['Table Grid'].font, BODY_FONT, BODY_SIZE)
    apply_font(doc.styles['List Bullet'].font, BODY_FONT, BODY_SIZE)


def _add_executive_summary(doc, df, risk_col, title):
    add_styled_heading(doc, title, 0)
    add_styled_heading(doc, '1. Executive Summary', 1)
    add_styled_paragraph(doc, f"Total Vulnerabilities Identified: {len(df)}")

    if risk_col in df.columns:
        risk_counts = df[risk_col].value_counts()
        ordered_risks = sorted(
            risk_counts.items(),
            key=lambda x: SEVERITY_ORDER.get(normalize_risk_key(x[0]), len(SEVERITY_ORDER))
        )
        add_styled_paragraph(doc, "Vulnerabilities by Risk Rating:")
        for risk_level, count in ordered_risks:
            add_styled_paragraph(doc, f"- {risk_level}: {count}", style='List Bullet')


def _add_bold_label(doc, text):
    add_styled_paragraph(doc, text, alignment=WD_ALIGN_PARAGRAPH.LEFT)
    for run in doc.paragraphs[-1].runs:
        apply_font(run.font, BODY_FONT, 10, bold=True)


def _add_finding_body(doc, row, affected_label, affected_col):
    _add_bold_label(doc, affected_label)
    add_text_block(doc, row.get(affected_col, 'N/A'))

    for heading_label, col_name, justify in FINDING_SECTIONS:
        _add_bold_label(doc, heading_label)
        add_text_block(doc, row.get(col_name, 'N/A'), justify=justify)


# --- WORD REPORT GENERATION FUNCTIONS ---

def create_web_vapt_report(df):
    doc = Document()
    _init_doc_styles(doc)
    risk_col = "CVSS Risk Rating" if "CVSS Risk Rating" in df.columns else "Risk Rating"
    _add_executive_summary(doc, df, risk_col, "Web Application VAPT Report")
    add_styled_heading(doc, '3. Detailed Findings', 1)
    df = sort_by_severity(df, risk_col)

    for index, row in df.iterrows():
        title = row.get('Issue Title', f'Unknown Finding #{index+1}')
        risk  = row.get(risk_col, 'Unknown')

        add_styled_heading(doc, f"3.{index+1} {title}", 2)
        add_five_col_summary_table(
            doc, risk,
            row.get('CVSS Score', 'N/A'),
            row.get('CVSS Vector', 'N/A'),
            row.get('OWASP Top 10', 'N/A'),
            row.get('CWE ID', 'N/A'),
            field4_label="OWASP Top 10",
            field5_label="CWE ID",
        )
        add_styled_paragraph(doc)
        _add_finding_body(doc, row, "Affected Module(s)", "Affected Module / URL")
        doc.add_page_break()

    buffer = BytesIO()
    doc.save(buffer)
    return buffer.getvalue()


def create_network_vapt_report(df):
    doc = Document()
    _init_doc_styles(doc)
    risk_col = "CVSS Risk Rating" if "CVSS Risk Rating" in df.columns else "Risk Rating"
    _add_executive_summary(doc, df, risk_col, "Network VAPT Report")
    add_styled_heading(doc, '3. Detailed Findings', 1)
    df = sort_by_severity(df, risk_col)

    for index, row in df.iterrows():
        title = row.get('Issue Title', f'Unknown Finding #{index+1}')
        risk  = row.get(risk_col, 'Unknown')

        add_styled_heading(doc, f"3.{index+1} {title}", 2)
        add_five_col_summary_table(
            doc, risk,
            row.get('CVSS Score', 'N/A'),
            row.get('CVSS Vector', 'N/A'),
            row.get('CVE ID', 'N/A'),
            row.get('Port/Protocol', 'N/A'),
            field4_label="CVE ID",
            field5_label="Port/Protocol",
        )
        add_styled_paragraph(doc)
        _add_finding_body(doc, row, "Affected Host/IP", "Affected Host/IP")
        doc.add_page_break()

    buffer = BytesIO()
    doc.save(buffer)
    return buffer.getvalue()


def create_host_config_report(df):
    doc = Document()
    _init_doc_styles(doc)
    risk_col = "CVSS Risk Rating" if "CVSS Risk Rating" in df.columns else "Risk Rating"
    _add_executive_summary(doc, df, risk_col, "Host Configuration Review Report")
    add_styled_heading(doc, '3. Detailed Findings', 1)
    df = sort_by_severity(df, risk_col)

    for index, row in df.iterrows():
        title = row.get('Issue Title', f'Unknown Finding #{index+1}')
        risk  = row.get(risk_col, 'Unknown')

        add_styled_heading(doc, f"3.{index+1} {title}", 2)
        add_host_config_summary_table(
            doc, risk,
            row.get('Benchmark/Standard', 'N/A'),
            row.get('Section/Control', 'N/A'),
        )
        add_styled_paragraph(doc)
        _add_finding_body(doc, row, "Affected Host/System", "Affected Host/System")
        doc.add_page_break()

    buffer = BytesIO()
    doc.save(buffer)
    return buffer.getvalue()


def create_cloud_vapt_report(df):
    doc = Document()
    _init_doc_styles(doc)
    risk_col = "CVSS Risk Rating" if "CVSS Risk Rating" in df.columns else "Risk Rating"
    _add_executive_summary(doc, df, risk_col, "Cloud VAPT Report")
    add_styled_heading(doc, '3. Detailed Findings', 1)
    df = sort_by_severity(df, risk_col)

    for index, row in df.iterrows():
        title = row.get('Issue Title', f'Unknown Finding #{index+1}')
        risk  = row.get(risk_col, 'Unknown')

        add_styled_heading(doc, f"3.{index+1} {title}", 2)
        add_five_col_summary_table(
            doc, risk,
            row.get('CVSS Score', 'N/A'),
            row.get('CVSS Vector', 'N/A'),
            row.get('Cloud Service', 'N/A'),
            row.get('Resource ID/ARN', 'N/A'),
            field4_label="Cloud Service",
            field5_label="Resource ID/ARN",
        )
        add_styled_paragraph(doc)
        _add_finding_body(doc, row, "Affected Cloud Resource", "Affected Cloud Resource")
        doc.add_page_break()

    buffer = BytesIO()
    doc.save(buffer)
    return buffer.getvalue()


REPORT_GENERATORS = {
    "Web Application VAPT":     create_web_vapt_report,
    "Network VAPT":              create_network_vapt_report,
    "Host Configuration Review": create_host_config_report,
    "Cloud VAPT":                create_cloud_vapt_report,
}


def create_excel_export(export_df, original_file_bytes):
    """Patch the original uploaded workbook in-place so all formatting,
    other sheets, and untouched cells are preserved. Only the cells that
    correspond to edited columns are overwritten."""
    wb = openpyxl.load_workbook(BytesIO(original_file_bytes))
    ws = wb["Risk Register"]

    # Build column-name → Excel column index from the header row (Excel row 2)
    col_map = {}
    for cell in ws[2]:
        if cell.value is not None:
            col_map[str(cell.value).strip()] = cell.column

    update_cols = [c for c in export_df.columns if c != '_excel_row']

    for _, row in export_df.iterrows():
        excel_row = int(row['_excel_row'])
        for col_name in update_cols:
            if col_name in col_map:
                ws.cell(row=excel_row, column=col_map[col_name], value=row[col_name])

    buffer = BytesIO()
    wb.save(buffer)
    return buffer.getvalue()


@st.cache_data
def load_data(file):
    df = pd.read_excel(file, sheet_name="Risk Register", header=1)
    if 'Issue Title' in df.columns:
        df = df.dropna(subset=['Issue Title'])
    # header=1 means data starts at Excel row 3; store so we can write back to the right cells
    df['_excel_row'] = df.index + 3
    return df


# --- STREAMLIT WEB DASHBOARD ---

st.set_page_config(page_title="VAPT Report Generator", layout="wide")

st.title("🛡️ Penetration Testing Report Generator")
st.write("Upload the VAPT Tracking List. The system will read the findings from the **Risk Register** tab and generate a dashboard and a downloadable Word Report.")
st.divider()

st.subheader("Step 1: Select Report Type")
report_type = st.radio(
    "Which type of report would you like to generate?",
    options=[
        "Network VAPT",
        "Web Application VAPT",
        "Host Configuration Review",
        "Cloud VAPT",
    ],
    index=None,
    horizontal=True,
)

if report_type is None:
    st.info("Please select a report type above to proceed.")
    st.stop()

st.success(f"Report type selected: **{report_type}**")
st.divider()

st.subheader("Step 2: Upload Excel File")
uploaded_file = st.file_uploader("Upload VAPT Excel file (.xlsx)", type=["xlsx", "xls"], key=f"uploader_{report_type}")

if uploaded_file is not None:
    try:
        df = load_data(uploaded_file)
        st.success("✅ File processed successfully!")

        risk_col = "CVSS Risk Rating" if "CVSS Risk Rating" in df.columns else "Risk Rating"
        df = sort_by_severity(df, risk_col)

        # --- EXECUTIVE SUMMARY DASHBOARD ---
        st.header("📊 Executive Summary")
        col1, col2 = st.columns([1, 2])
        with col1:
            st.metric("Total Vulnerabilities Identified", len(df))
        with col2:
            if risk_col in df.columns:
                risk_counts = df[risk_col].value_counts()
                chart_df = pd.DataFrame({
                    "Risk Rating": risk_counts.index.astype(str),
                    "Count": risk_counts.values,
                })
                chart_df["Color"] = chart_df["Risk Rating"].apply(
                    lambda r: f"#{get_risk_style(r)['bg']}"
                )
                chart_df = chart_df.assign(
                    _order=chart_df["Risk Rating"].apply(
                        lambda r: SEVERITY_ORDER.get(normalize_risk_key(r), len(SEVERITY_ORDER))
                    )
                ).sort_values("_order").drop(columns="_order")

                color_scale = alt.Scale(
                    domain=chart_df["Risk Rating"].tolist(),
                    range=chart_df["Color"].tolist(),
                )
                max_count = int(chart_df["Count"].max())
                chart = (
                    alt.Chart(chart_df)
                    .mark_bar()
                    .encode(
                        x=alt.X("Risk Rating:N", sort=chart_df["Risk Rating"].tolist(), title="Risk Rating"),
                        y=alt.Y("Count:Q", title="Count",
                                scale=alt.Scale(domain=[0, max_count], nice=False),
                                axis=alt.Axis(tickMinStep=1, format='d', tickCount=max_count)),
                        color=alt.Color("Risk Rating:N", scale=color_scale, legend=None),
                        tooltip=["Risk Rating", "Count"],
                    )
                )
                st.altair_chart(chart, width='stretch')

        st.divider()
        st.header("📝 Detailed Findings")
        st.caption("All fields are editable — your changes will be reflected in the downloaded Word document.")

        # Reset edit state when the file or report type changes
        file_id = f"{report_type}_{uploaded_file.name}_{uploaded_file.size}"
        if st.session_state.get('_file_id') != file_id:
            st.session_state['_file_id'] = file_id
            for k in [k for k in list(st.session_state.keys()) if k.startswith('edit_')]:
                del st.session_state[k]
            st.session_state.pop('saved_word_bytes', None)
            st.session_state.pop('saved_excel_bytes', None)

        def _val(v):
            if v is None:
                return ''
            try:
                if pd.isna(v):
                    return ''
            except (TypeError, ValueError):
                pass
            return str(v)

        meta_fields = REPORT_EDIT_FIELDS[report_type]
        affected_label, affected_col = AFFECTED_ASSET_CONFIG[report_type]

        # Pre-populate session state from the DataFrame on first load
        for i, (idx, row) in enumerate(df.iterrows()):
            defaults = {
                f"edit_title_{i}":    _val(row.get('Issue Title')),
                f"edit_risk_{i}":     _val(row.get(risk_col)),
                f"edit_status_{i}":   _val(row.get('Status')) or 'Open',
                f"edit_affected_{i}": _val(row.get(affected_col)),
                f"edit_obs_{i}":      _val(row.get('Observations')),
                f"edit_impl_{i}":     _val(row.get('Implications')),
                f"edit_rec_{i}":      _val(row.get('Recommendations')),
                f"edit_mgmt_{i}":     _val(row.get('Management Comments')),
                f"edit_followup_{i}": _val(row.get('Post Review Observations')),
            }
            for label, key_prefix, col_name in meta_fields:
                defaults[f"{key_prefix}_{i}"] = _val(row.get(col_name))
            for k, v in defaults.items():
                if k not in st.session_state:
                    st.session_state[k] = v

        for i, (idx, row) in enumerate(df.iterrows()):
            label_title = st.session_state.get(f"edit_title_{i}") or f"Finding #{i+1}"
            label_risk  = st.session_state.get(f"edit_risk_{i}", "")
            with st.expander(f"3.{i+1} [{label_risk.upper()}] {label_title}"):
                c1, c2, c3 = st.columns(3)
                with c1:
                    st.text_input("Issue Title", key=f"edit_title_{i}")
                    st.text_input("Risk Rating", key=f"edit_risk_{i}")
                    st.text_input("Status", key=f"edit_status_{i}")
                with c2:
                    for label, key_prefix, _ in meta_fields[:2]:
                        st.text_input(label, key=f"{key_prefix}_{i}")
                with c3:
                    for label, key_prefix, _ in meta_fields[2:]:
                        st.text_input(label, key=f"{key_prefix}_{i}")
                st.text_area(affected_label, key=f"edit_affected_{i}", height=80)
                st.text_area("Observation", key=f"edit_obs_{i}", height=150)
                st.text_area("Implication", key=f"edit_impl_{i}", height=150)
                st.text_area("Recommendations", key=f"edit_rec_{i}", height=150)
                st.text_area("Management Comments", key=f"edit_mgmt_{i}", height=100)
                st.text_area("Follow-up Comments", key=f"edit_followup_{i}", height=100)

        # --- SAVE & DOWNLOAD ---
        st.divider()
        st.header("📥 Export Report")

        st.markdown("""
        <style>
        div[data-testid="stButton"] > button,
        div[data-testid="stDownloadButton"] > button {
            height: 2.75rem;
            min-height: 2.75rem;
        }
        </style>
        """, unsafe_allow_html=True)

        col_save, col_dl, col_xl, _ = st.columns([1, 1, 1, 3])

        with col_save:
            if st.button("💾 Save Changes", type="primary", use_container_width=True):
                export_df = df.copy().astype(object)
                for i, idx in enumerate(df.index):
                    export_df.at[idx, 'Issue Title']              = st.session_state.get(f"edit_title_{i}", '')
                    export_df.at[idx, risk_col]                   = st.session_state.get(f"edit_risk_{i}", '')
                    export_df.at[idx, 'Status']                   = st.session_state.get(f"edit_status_{i}", '')
                    export_df.at[idx, affected_col]               = st.session_state.get(f"edit_affected_{i}", '')
                    export_df.at[idx, 'Observations']             = st.session_state.get(f"edit_obs_{i}", '')
                    export_df.at[idx, 'Implications']             = st.session_state.get(f"edit_impl_{i}", '')
                    export_df.at[idx, 'Recommendations']          = st.session_state.get(f"edit_rec_{i}", '')
                    export_df.at[idx, 'Management Comments']      = st.session_state.get(f"edit_mgmt_{i}", '')
                    export_df.at[idx, 'Post Review Observations'] = st.session_state.get(f"edit_followup_{i}", '')
                    for label, key_prefix, col_name in meta_fields:
                        export_df.at[idx, col_name] = st.session_state.get(f"{key_prefix}_{i}", '')

                generate_fn = REPORT_GENERATORS[report_type]
                st.session_state['saved_word_bytes'] = generate_fn(export_df)
                st.session_state['saved_excel_bytes'] = create_excel_export(export_df, uploaded_file.getvalue())

        with col_dl:
            has_bytes = 'saved_word_bytes' in st.session_state
            report_filename = report_type.replace(" ", "_") + "_Report.docx"
            st.download_button(
                label="📄 Download Report (.docx)",
                data=st.session_state.get('saved_word_bytes', b''),
                file_name=report_filename,
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                disabled=not has_bytes,
                use_container_width=True,
            )

        with col_xl:
            excel_filename = report_type.replace(" ", "_") + "_Updated.xlsx"
            st.download_button(
                label="📊 Download Excel (.xlsx)",
                data=st.session_state.get('saved_excel_bytes', b''),
                file_name=excel_filename,
                mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
                disabled='saved_excel_bytes' not in st.session_state,
                use_container_width=True,
            )

        if 'saved_word_bytes' in st.session_state:
            st.success("✅ Changes saved! Download the Word report or the updated Excel file.")

    except Exception as e:
        st.error(f"An error occurred: {e}")
        st.info("Please ensure the uploaded Excel file contains a tab named exactly 'Risk Register'.")
